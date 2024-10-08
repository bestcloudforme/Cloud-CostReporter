import boto3
import re
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, RGBColor
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime, timedelta

AMOUNT_OF_MONTHS = 2
THRESHOLD_PERCENTAGE = 5
THRESHOLD_USD = 0.5
MINIMUM_COST_PER_SERVICE = 1 # first month + second month

# AWS Cost Explorer client
client = boto3.client('ce')


# Helper function to get the start and end dates for the past X months
def get_date_range(months):
    end = datetime.today().replace(day=1)  # First day of current month
    start = (end - timedelta(days=months * 30)).replace(day=1)  # First day X months ago
    return start.strftime('%Y-%m-%d'), end.strftime('%Y-%m-%d')

# Get cost data for the last 3 months, grouped by service
def fetch_cost_data():
    start_date, end_date = get_date_range(AMOUNT_OF_MONTHS)
    response = client.get_cost_and_usage(
        TimePeriod={'Start': start_date, 'End': end_date},
        Granularity='MONTHLY',
        Metrics=['UnblendedCost'],
        GroupBy=[{'Type': 'DIMENSION', 'Key': 'SERVICE'}]
    )
    return response['ResultsByTime']

def fetch_cost_data_by_service_name(service_name, group_by='CostCenter'):
    if group_by == 'CostCenter':
        group_key = {'Type': 'TAG', 'Key': 'CostCenter'}
    elif group_by == 'UsageType':
        group_key = {'Type': 'DIMENSION', 'Key': 'USAGE_TYPE'}
    else:
        raise ValueError("Invalid group_by value. Must be 'CostCenter' or 'UsageType'.")

    start_date, end_date = get_date_range(AMOUNT_OF_MONTHS)
    
    response = client.get_cost_and_usage(
        TimePeriod={'Start': start_date, 'End': end_date},
        Granularity='MONTHLY',
        Metrics=['UnblendedCost'],
        GroupBy=[group_key],  # Group by the specified key (CostCenter or UsageType)
        Filter={
            'Dimensions': {
                'Key': 'SERVICE',  # Filter by the service name
                'Values': [service_name]  # Service name provided by the function parameter
            }
        }
    )
    
    return response['ResultsByTime']

# Process cost data for plotting and comparison
def process_cost_data(data):
    service_costs = {}
    monthly_totals = {}

    for month_data in data:
        month = month_data['TimePeriod']['Start']
        monthly_total = 0

        for group in month_data['Groups']:
            service = group['Keys'][0]
            cost = float(group['Metrics']['UnblendedCost']['Amount'])
            monthly_total += cost

            if service not in service_costs:
                service_costs[service] = [0] * len(data)  # Initialize list for each month
            service_costs[service][data.index(month_data)] = cost

        monthly_totals[month] = monthly_total

    return service_costs, monthly_totals

# Plot total costs for top 9 services
def plot_cost_graph(service_costs, monthly_totals):
    sorted_services = sorted(service_costs.items(), key=lambda x: sum(x[1]), reverse=True)
    top_services = dict(sorted_services[:9])  # Get top 9 services
    other_services_total = [sum(x) for x in zip(*[costs for service, costs in sorted_services[9:]])]

    # Colors for the graph
    colors = plt.colormaps['tab20'](range(len(top_services) + 1))

    # Plot the total costs for top 9 services
    plt.figure(figsize=(12, 6))
    bottom = [0] * len(monthly_totals)
    months = list(monthly_totals.keys())

    for i, (service, costs) in enumerate(top_services.items()):
        plt.bar(months, costs, bottom=bottom, color=colors[i], label=service)
        bottom = [a + b for a, b in zip(bottom, costs)]

    # Plot "Others" category
    if other_services_total:
        plt.bar(months, other_services_total, bottom=bottom, color=colors[len(top_services)], label='Others')

    plt.xlabel('Month')
    plt.ylabel('Total Cost (USD)')
    plt.title('Monthly Total Costs')
    plt.legend(loc='upper center', bbox_to_anchor=(0.5, -0.1), ncol=3)
    plt.grid(True)
    plt.tight_layout()
    plt.savefig('monthly_total_costs_fixed.png')
    plt.close()

def add_total_cost_comparison_table(doc, monthly_totals):
    # Get the last three months for comparison
    months = [datetime.strptime(month, '%Y-%m-%d').strftime('%B') for month in sorted(monthly_totals.keys())[-AMOUNT_OF_MONTHS:]]
    
    # Prepare the table structure
    table = doc.add_table(rows=1, cols=len(months) + 2)
    table.style = 'Normal Table'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Total Cost'
    for i, month in enumerate(months):
        hdr_cells[i + 1].text = month
    hdr_cells[len(months) + 1].text = 'Difference (USD)'

    # Calculate the total costs for each month
    total_costs = [monthly_totals[month] for month in sorted(monthly_totals.keys())[-AMOUNT_OF_MONTHS:]]
    
    # Calculate the difference between the last two months
    cost_diff = total_costs[-1] - total_costs[-2]

    # Add total costs row to the table
    row_cells = table.add_row().cells
    row_cells[0].text = 'Total'
    for i, cost in enumerate(total_costs):
        row_cells[i + 1].text = f'{cost:.2f}'
    
    # Add the difference column value
    diff_cell = row_cells[len(months) + 1]
    diff_cell.text = f'{cost_diff:.2f}'
    
    # Set the color based on the difference
    if cost_diff > 0:
        # Set text color to red for increased cost
        for run in diff_cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red
    elif cost_diff < 0:
        # Set text color to green for decreased cost
        for run in diff_cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0, 255, 0)  # Green

def add_cost_comparison_table(doc, service_costs, monthly_totals):
    # Get the last two months for comparison
    last_two_months = [datetime.strptime(month, '%Y-%m-%d').strftime('%B') for month in sorted(monthly_totals.keys())[-2:]]
    
    # Prepare the table structure
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Medium Shading 1 Accent 6'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Service'
    hdr_cells[1].text = f'{last_two_months[0]} (USD)'
    hdr_cells[2].text = f'{last_two_months[1]} (USD)'
    hdr_cells[3].text = 'Difference (USD)'

    # Create a list to store service differences
    service_diff = []
    
    # Iterate through services to compare the last two months
    for service, costs in service_costs.items():
        # Extract the cost for the last two months
        cost_prev = costs[-2]  # Cost for second to last month
        cost_last = costs[-1]  # Cost for the last month
        
        # Calculate the difference
        difference = cost_last - cost_prev
        
        # Filter out services where the last month's cost is below $50
        # if cost_last >= 50:
            # Append the service and its costs to the list
        service_diff.append((service, cost_prev, cost_last, difference))

    # Sort the services based on the difference
    service_diff.sort(key=lambda x: x[3], reverse=True)
    
    # Add services to the table
    for service, cost_prev, cost_last, difference in service_diff:
        row_cells = table.add_row().cells
        row_cells[0].text = service
        row_cells[1].text = f'{cost_prev:.2f}'
        row_cells[2].text = f'{cost_last:.2f}'
        
        # Set the color based on the difference
        diff_cell = row_cells[3]
        diff_cell.text = f'{difference:.2f}'
        if difference > 0:
            # Set text color to red for increased cost
            for run in diff_cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        elif difference < 0:
            # Set text color to green for decreased cost
            for run in diff_cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0, 255, 0)


def set_cell_background(cell, color):
    """Helper function to set background color of a table cell."""
    # Get the table cell properties element (tcPr)
    tcPr = cell._element.get_or_add_tcPr()
    # Create and append the color shading element
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)  # Set the fill color
    tcPr.append(shd)

def add_service_compare_tables(doc, service_compares):
    """
    Adds service comparison data to the DOCX document as tables, for both CostCenter and UsageType,
    with improved styling and a light teal background for headers.

    Args:
        doc: The Document object where the tables will be added.
        service_compares (dict): A dictionary with separate keys for CostCenter and UsageType comparisons.
    """
    for service_name, comparison_data in service_compares.items():
        doc.add_heading(service_name, level=2)

        # CostCenter table
        if comparison_data.get("CostCenter"):
            doc.add_heading(f'{service_name} - Cost Center Comparison', level=3)
            cost_center_table = doc.add_table(rows=1, cols=4)
            cost_center_table.style = 'Table Grid'  # Apply a clean table style
            hdr_cells = cost_center_table.rows[0].cells

            # Set headers, apply bold formatting, and set background color
            hdr_cells[0].text = 'Cost Center'
            hdr_cells[1].text = 'First Month Cost (USD)'
            hdr_cells[2].text = 'Second Month Cost (USD)'
            hdr_cells[3].text = 'Difference (USD)'

            for cell in hdr_cells:
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                cell.paragraphs[0].alignment = 1  # Center-align header text
                set_cell_background(cell, 'CCFFFF')  # Light teal background

            # Add data rows
            for cost_center, first_month_cost, second_month_cost, difference in comparison_data["CostCenter"]:
                row_cells = cost_center_table.add_row().cells
                row_cells[0].text = str(cost_center)
                row_cells[1].text = f"{first_month_cost:.2f}"
                row_cells[2].text = f"{second_month_cost:.2f}"
                row_cells[3].text = f"{difference:.2f}"

                # Center-align numeric columns
                for i in range(1, 4):
                    row_cells[i].paragraphs[0].alignment = 1  # Center-align

        # UsageType table
        if comparison_data.get("UsageType"):
            doc.add_heading(f'{service_name} - Usage Type Comparison', level=3)
            usage_type_table = doc.add_table(rows=1, cols=4)
            usage_type_table.style = 'Table Grid'  # Apply a clean table style
            hdr_cells = usage_type_table.rows[0].cells

            # Set headers, apply bold formatting, and set background color
            hdr_cells[0].text = 'Usage Type'
            hdr_cells[1].text = 'First Month Cost (USD)'
            hdr_cells[2].text = 'Second Month Cost (USD)'
            hdr_cells[3].text = 'Difference (USD)'

            for cell in hdr_cells:
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                cell.paragraphs[0].alignment = 1  # Center-align header text
                set_cell_background(cell, 'FFFFFF')  # Light teal background

            # Add data rows
            for usage_type, first_month_cost, second_month_cost, difference in comparison_data["UsageType"]:
                row_cells = usage_type_table.add_row().cells
                row_cells[0].text = str(usage_type)
                row_cells[1].text = f"{first_month_cost:.2f}"
                row_cells[2].text = f"{second_month_cost:.2f}"
                row_cells[3].text = f"{difference:.2f}"
                # Center-align numeric columns
                for i in range(1, 4):
                    row_cells[i].paragraphs[0].alignment = 1  # Center-align

    return doc





def generate_service_compares(service_names):
    service_compares = {}

    for service_name in service_names:
        cost_center_cost = fetch_cost_data_by_service_name(service_name, group_by="CostCenter")

        usage_type_cost = fetch_cost_data_by_service_name(service_name, group_by="UsageType")

        cost_center_data = []
        groups1 = {group['Keys'][0]: group for group in cost_center_cost[0]['Groups']}
        groups2 = {group['Keys'][0]: group for group in cost_center_cost[1]['Groups']}
        all_keys = set(groups1.keys()).union(set(groups2.keys()))

        for key in all_keys:
            cost_center = re.sub(r'^CostCenter\$$', 'No tag key: CostCenter', key)
            first_month_cost = float(groups1.get(key, {'Metrics': {'UnblendedCost': {'Amount': '0'}}})['Metrics']['UnblendedCost']['Amount'])
            second_month_cost = float(groups2.get(key, {'Metrics': {'UnblendedCost': {'Amount': '0'}}})['Metrics']['UnblendedCost']['Amount'])
            difference = second_month_cost - first_month_cost
            total_cost = first_month_cost + second_month_cost

            if total_cost >= MINIMUM_COST_PER_SERVICE:
                cost_center_data.append([cost_center, first_month_cost, second_month_cost, difference])

        cost_center_data.sort(key=lambda x: x[2], reverse=True)

        usage_type_data = []
        groups1 = {group['Keys'][0]: group for group in usage_type_cost[0]['Groups']}
        groups2 = {group['Keys'][0]: group for group in usage_type_cost[1]['Groups']}
        all_keys = set(groups1.keys()).union(set(groups2.keys()))

        for key in all_keys:
            first_month_cost = float(groups1.get(key, {'Metrics': {'UnblendedCost': {'Amount': '0'}}})['Metrics']['UnblendedCost']['Amount'])
            second_month_cost = float(groups2.get(key, {'Metrics': {'UnblendedCost': {'Amount': '0'}}})['Metrics']['UnblendedCost']['Amount'])
            difference = second_month_cost - first_month_cost
            total_cost = first_month_cost + second_month_cost

            if total_cost >= MINIMUM_COST_PER_SERVICE:
                usage_type_data.append([key, first_month_cost, second_month_cost, difference])

        usage_type_data.sort(key=lambda x: x[2], reverse=True)

        service_compares[service_name] = {
            "CostCenter": cost_center_data,
            "UsageType": usage_type_data
        }

    return service_compares



def calculate_monthly_difference_from_percentage(
    service_costs,
    percentage_threshold: int,
    usd_threshold: int
):
    monthly_diff = {}
    for service_name, cost in service_costs.items():
        difference = cost[1] - cost[0]
        difference_percentage = (difference) / (cost[0] + 0.000000001) * 100 #TODO: fix

        if difference > usd_threshold and difference_percentage > percentage_threshold:
            monthly_diff[service_name] = difference

    return dict(sorted(monthly_diff.items(), key=lambda x: x[1], reverse=True))

# Generate the Word document
def generate_report(service_costs, monthly_totals, service_compares):
    doc = Document()
    doc.add_heading('AWS Cost Analysis', 0)
    
    # Add the total cost graph
    doc.add_heading('Total Service Cost Graph', level=1)
    doc.add_picture('monthly_total_costs_fixed.png', width=Inches(6.0))
    doc.add_paragraph('Monthly Total Costs for the last 3 months.')
    
    # Add the total cost comparison table
    doc.add_heading('Total Cost Comparison', level=1)
    add_total_cost_comparison_table(doc, monthly_totals)

    # Add the cost comparison table
    doc.add_heading('Top Service Cost Comparison', level=1)
    add_cost_comparison_table(doc, service_costs, monthly_totals)

    doc.add_heading(f"Cost analysis for services that increased more than {THRESHOLD_PERCENTAGE}")
    add_service_compare_tables(doc, service_compares)

    # Save the document
    doc.save('AWS_Cost_Analysis_Report.docx')

# Main function to coordinate everything
def main():
    data = fetch_cost_data()
    service_costs, monthly_totals = process_cost_data(data)
    monthly_diff = calculate_monthly_difference_from_percentage(service_costs,
                                                                percentage_threshold=THRESHOLD_PERCENTAGE,
                                                                usd_threshold=THRESHOLD_USD)

    service_compares = generate_service_compares(monthly_diff.keys())
    plot_cost_graph(service_costs, monthly_totals)
    generate_report(service_costs, monthly_totals, service_compares)

if __name__ == "__main__":
    main()
